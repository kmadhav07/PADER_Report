"""
Professional DOCX Exporter Module (python-docx Table Parser).

Author: Madhav Kumar
Module: exporters.docx_exporter
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_background(cell, fill_color):
    """Set background color of a Word table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)


def parse_markdown_table_docx(table_lines: list, doc: Document):
    """Parse Markdown table lines into native Word table."""
    raw_rows = []
    for line in table_lines:
        if '| ---' in line or '|:---' in line:
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells and any(c for c in cells):
            raw_rows.append(cells)

    if not raw_rows:
        return

    num_rows = len(raw_rows)
    num_cols = max(len(r) for r in raw_rows)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'

    for row_idx, row in enumerate(raw_rows):
        is_header = (row_idx == 0)
        for col_idx, cell_text in enumerate(row):
            if col_idx < num_cols:
                cell = table.cell(row_idx, col_idx)
                cell.text = cell_text.replace('**', '').replace('*', '')
                
                if is_header:
                    set_cell_background(cell, '1E3A8A')
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.name = 'Arial'
                            run.font.size = Pt(9.5)
                else:
                    if row_idx % 2 == 1:
                        set_cell_background(cell, 'F8FAFC')
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(9)


def export_to_docx(markdown_text: str, output_path: str) -> str:
    """Export report to Word DOCX format using python-docx."""
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    doc = Document()
    lines = markdown_text.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        if not line:
            i += 1
            continue

        # Table detection
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            parse_markdown_table_docx(table_lines, doc)
            continue

        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### ') or line.startswith('##### '):
            doc.add_heading(line.lstrip('#').strip(), level=4)
        elif line == '---':
            pass
        else:
            p_text = line.replace('**', '').replace('*', '')
            doc.add_paragraph(p_text)

        i += 1

    doc.save(output_path)
    return output_path
